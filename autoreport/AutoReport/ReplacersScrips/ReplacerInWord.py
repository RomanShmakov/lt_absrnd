import math
import os
import shutil
import zipfile
import yaml
from PIL import Image
from docx import Document

from ReplacersScrips.AlfagenAPIDescriptionGraf import get_description_for_graf, get_description_for_table
from SaverScripts.GrafanaAPIValueSaver import get_content_from_panel
import pandas as pd
import json

# config #
from constants import (CONFIG_NAME)

# config #
base_path = '.'
config_file = os.path.join(base_path, CONFIG_NAME)
config = yaml.safe_load(open(config_file, 'r'))


# *Заменяем*слова*в*таблице*с*результатами******************************************************************************
def get_avg_values_from_panel(panel_name, time_from, time_to):
    res = get_content_from_panel(panel_name, time_from, time_to)

    average_value = None

    if res is not None or not res:
        values = []
        # Извлечение данных по пути: $.results.A.frames[0].data.values[1][*]
        try:
            # total_count = sum(frame['data']['values'][1])
            for entry in res['results']['A']['frames'][0]['data']['values'][1]:
                values.append(float(entry))
        except (KeyError, IndexError) as e:
            print(f"Ошибка при извлечении данных: {e}")
            values = []
        # Вычисляем среднее значение
        if values:
            average_value = sum(values) / len(values)
            print(f"Панель: {panel_name}, среднее значение до округления: {average_value}")
            average_value = math.ceil(average_value)
            print(f"Панель: {panel_name}, среднее значение после округления: {average_value}")
        else:
            print("Нет данных для вычисления среднего значения")
    else:
        print(f"Ошибка при получении контента из панели: {panel_name}")

    return average_value


# *Заменяем*слова*в*таблице*с*результатами******************************************************************************

# *Функция*для*замены*текста*в*документе********************************************************************************
def repace_table_result(doc_name):
    time_from = config["target_step_start_time"]
    time_to = config["target_step_end_time"]

    throughput = get_avg_values_from_panel("Throughput", time_from, time_to)
    response_time = get_avg_values_from_panel("Response Time - Pass (95th pct)", time_from, time_to)
    percent_error = get_avg_values_from_panel("% Error (< 20%)", time_from, time_to)
    threads = get_avg_values_from_panel("Active Threads", time_from, time_to)

    doc = Document(doc_name)
    table = doc.tables[1]

    # получаем все объекты с ячейками в строке для вставки
    cells = table.rows[1].cells
    cells[0].text = str(f"{throughput} rps")
    cells[1].text = str(f"{response_time} ms")
    cells[2].text = str(f"{percent_error} %")
    cells[3].text = str(f"{threads} VU")

    doc.save(doc_name)
    print(f"В документе {doc_name} сохранена таблица с результатами теста")


def replace_in_paragraph(paragraph, old, new):
    """Заменяем old->new в параграфе, даже если old разбит по runs."""
    if old not in paragraph.text:
        return False
    replaced_text = paragraph.text.replace(old, new)
    # удалить все runs
    for run in paragraph.runs:
        run.text = ''
    # добавить один run с новым текстом
    paragraph.add_run(replaced_text)
    return True


def replace_in_table(table, old, new):
    replaced = False
    for row in table.rows:
        for cell in row.cells:
            # у ячейки могут быть и таблицы вложенные
            for paragraph in cell.paragraphs:
                if replace_in_paragraph(paragraph, old, new):
                    replaced = True
            for inner_table in cell.tables:
                if replace_in_table(inner_table, old, new):
                    replaced = True
    return replaced


def replace_in_doc(doc, old_text, new_text):
    old = old_text
    new = str(new_text)  # обязательно строка
    was_replaced = False

    # параграфы основного документа
    for paragraph in doc.paragraphs:
        if replace_in_paragraph(paragraph, old, new):
            was_replaced = True

    # таблицы в основном документе
    for table in doc.tables:
        if replace_in_table(table, old, new):
            was_replaced = True

    # заголовки/футеры секций
    for section in doc.sections:
        header = section.header
        footer = section.footer
        # header.paragraphs + header.tables
        for paragraph in header.paragraphs:
            if replace_in_paragraph(paragraph, old, new):
                was_replaced = True
        for table in header.tables:
            if replace_in_table(table, old, new):
                was_replaced = True
        # footer
        for paragraph in footer.paragraphs:
            if replace_in_paragraph(paragraph, old, new):
                was_replaced = True
        for table in footer.tables:
            if replace_in_table(table, old, new):
                was_replaced = True

    return was_replaced


def replace_text(doc_name, old_text, new_text):
    doc = Document(doc_name)
    was_replase = replace_in_doc(doc, old_text, new_text)
    doc.save(doc_name)
    if was_replase:
        print("Text was replaced success")
    else:
        print("Error: text not found, replace failed")


# *Функция*для*замены*текста*в*документе********************************************************************************


# **Заменяем*описания*для*графиков*ответом*альфагена********************************************************************


def replace_description_for_graf(doc_name):
    typeTest = config["typeTest"]

    if typeTest == "LoadTest":
        purpose_test = "Это тест Поиск максимальной производительности системы. На систему подаётся нагрузка, генерируемая виртуальными пользователями через JMeter на приложение, которое работает с YDB (yandex database). " \
                       "Нагрузка постепенно повышается ступенями, цель теста найти пик нагрузки, который может выдержать система. После выхода на новую ступень нагрузки, мы немного ждем, чтобы зафиксировать результаты."
    elif typeTest == "StressTest":
        purpose_test = "Это Стресс тест. На систему подаётся нагрузка, генерируемая виртуальными пользователями через JMeter на приложение, которое работает с YDB (yandex database). " \
                       "Сначала мы подаем стабильную нагрузку, которую без проблем может выдержать система, а после на небольшой промежуток времени значительно увеличиваем нагрузку " \
                       "и фиксируем поведение системы сверх своих возможностей, чтобы увидеть как система ломается и как восстанавливается"
    elif typeTest == "VolumeTest":
        purpose_test = "Это тест объема. После значительного увеличения объема данных БД, на систему подаётся стабильная нагрузка в течение долгого времени, генерируемая виртуальными пользователями через JMeter на приложение, которое работает с YDB (yandex database)." \
                       "Цель теста увидеть как система справляется с обработкой очень большого количества данных, а не пользователей"
    elif typeTest == "StabilityTest":
        purpose_test = "Это тест стабильности. На систему подаётся стабильная нагрузка в течение долгого времени, генерируемая виртуальными пользователями через JMeter на приложение, которое работает с YDB (yandex database)."
    elif typeTest == "FailoverTest":
        purpose_test = "Это тест отказоустойчивости. На систему подаётся нагрузка, генерируемая виртуальными пользователями через JMeter на приложение, которое работает с YDB (yandex database)."
    elif typeTest == "ScalabilityTest":
        purpose_test = "Это тест масштабируемости. На систему подаётся стабильная нагрузка в течение долгого времени, генерируемая виртуальными пользователями через JMeter на приложение, которое работает с YDB (yandex database). " \
                       "Сначала мы подаем стабильную нагрузку, а после добавляем ей ресурсы и смотрим, помогло ли это улучшить производительность системы."
    else:
        purpose_test = "Error to get typeTest from config"

    # JMETER #

    # 1_Throughput
    image_name = 'save/image1.png'
    purpose_graf = "Это график JMeter - Throughput, он показывает интенсивности запросов в секунду."
    new_description = get_description_for_graf(image_name, purpose_graf, purpose_test)
    replace_text(doc_name, "<LLM_Throughput>", new_description)

    # 2_VirtualUsers
    image_name = 'save/image2.png'
    purpose_graf = "Это график JMeter - VirtualUsers, он показывает выход пользователей на рабочую нагрузку."
    new_description = get_description_for_graf(image_name, purpose_graf, purpose_test)
    replace_text(doc_name, "<LLM_VirtualUsers>", new_description)

    # 3_ResponseTime
    image_name = 'save/image3.png'
    purpose_graf = "Это график JMeter - ResponseTime, он показывает время ответа по успешным запросам."
    new_description = get_description_for_graf(image_name, purpose_graf, purpose_test)
    replace_text(doc_name, "<LLM_ResponseTime>", new_description)

    # 4_PercentErrors
    image_name = 'save/image4.png'
    purpose_graf = "Это график JMeter - PercentErrors, он показывает"
    new_description = get_description_for_graf(image_name, purpose_graf, purpose_test)
    replace_text(doc_name, "<LLM_PercentErrors>", new_description)

    # Count error in table

    doc = Document(doc_name)
    table = doc.tables[2]
    table_error_data = []
    for row in range(1, 6):
        # получаем все объекты с ячейками в строке для вставки
        cells = table.rows[row].cells
        table_error_data.append(cells[2].text)
        table_error_data.append(cells[3].text)
    string_with_all_errors = ", ".join(table_error_data)
    purpose_graf = "Это таблица JMeter - CountErrorInTable, она показывает количество ошибок за весь тест по каждому сообщению."
    new_description = get_description_for_table(string_with_all_errors, purpose_graf, purpose_test)
    replace_text(doc_name, "<LLM_CountErrorInTable>", new_description)

    # 5_ErrorByMessage
    image_name = 'save/image5.png'
    purpose_graf = "Это график JMeter - ErrorByMessage, он показывает количество ошибок (за период агрегации 5 секунд) в разрезе времени."
    new_description = get_description_for_graf(image_name, purpose_graf, purpose_test)
    replace_text(doc_name, "<LLM_ErrorByMessage>", new_description)

    # DB status #

    # 6_QueryLatency
    image_name = 'save/image6.png'
    purpose_graf = "Это график YDB - DB status - QueryLatency, он показывает время выполнения запроса от момента получения запроса YDB до возврата результата."
    new_description = get_description_for_graf(image_name, purpose_graf, purpose_test)
    replace_text(doc_name, "<LLM_QueryLatency>", new_description)

    # 7_SessionCount
    image_name = 'save/image7.png'
    purpose_graf = "Это график YDB - DB status - SessionCount, он показывает количество сессий БД на протяжении всего теста."
    new_description = get_description_for_graf(image_name, purpose_graf, purpose_test)
    replace_text(doc_name, "<LLM_SessionCount>", new_description)

    # DB overview #

    # 8_Requests
    image_name = 'save/image8.png'
    purpose_graf = "Это график YDB - DB Overview - Requests, он показывает количество выполненных запросов в секунду "
    new_description = get_description_for_graf(image_name, purpose_graf, purpose_test)
    replace_text(doc_name, "<LLM_Requests>", new_description)

    # 9_ErrorsType
    image_name = 'save/image9.png'
    purpose_graf = "Это график YDB - DB Overview - ErrorsType, он показывает интенсивность неуспешных запросов в секунду в разрезе типа ошибки."
    new_description = get_description_for_graf(image_name, purpose_graf, purpose_test)
    replace_text(doc_name, "<LLM_ErrorsType>", new_description)

    # 10_DroppedRequests
    image_name = 'save/image10.png'
    purpose_graf = "Это график YDB - DB Overview - DroppedRequests, он показывает отброшенные запросы БД."
    new_description = get_description_for_graf(image_name, purpose_graf, purpose_test)
    replace_text(doc_name, "<LLM_DroppedRequests>", new_description)

    # 11_DroppedResponse
    image_name = 'save/image11.png'
    purpose_graf = "Это график YDB - DB Overview - DroppedResponse, он показывает отброшенные ответы БД."
    new_description = get_description_for_graf(image_name, purpose_graf, purpose_test)
    replace_text(doc_name, "<LLM_DroppedResponse>", new_description)

    # 12_RequestsInFlight
    image_name = 'save/image12.png'
    purpose_graf = "Это график YDB - DB Overview - RequestsInFlight, он показывает количество запросов в обработке БД на протяжении всего теста."
    new_description = get_description_for_graf(image_name, purpose_graf, purpose_test)
    replace_text(doc_name, "<LLM_RequestsInFlight>", new_description)

    # 13_UserPool
    image_name = 'save/image13.png'
    purpose_graf = "Это график YDB - DB Overview - UserPool, он показывает потребление CPU пулом для пользовательских операций."
    new_description = get_description_for_graf(image_name, purpose_graf, purpose_test)
    replace_text(doc_name, "<LLM_UserPool>", new_description)

    # 14_SystemPool
    image_name = 'save/image14.png'
    purpose_graf = "Это график YDB - DB Overview - SystemPool, он показывает потребление CPU пулом для системных задач."
    new_description = get_description_for_graf(image_name, purpose_graf, purpose_test)
    replace_text(doc_name, "<LLM_SystemPool>", new_description)

    # 15_ICPool
    image_name = 'save/image15.png'
    purpose_graf = "Это график YDB - DB Overview - ICPool, он показывает потребление CPU пулом межсетевой коммуникации (Interconnect Pool) между узлами БД."
    new_description = get_description_for_graf(image_name, purpose_graf, purpose_test)
    replace_text(doc_name, "<LLM_ICPool>", new_description)

    # 16_BatchPool
    image_name = 'save/image16.png'
    purpose_graf = "Это график YDB - DB Overview - BatchPool, он показывает потребление CPU пулом пакетной обработки."
    new_description = get_description_for_graf(image_name, purpose_graf, purpose_test)
    replace_text(doc_name, "<LLM_BatchPool>", new_description)

    # 17_IOPool
    image_name = 'save/image17.png'
    purpose_graf = "Это график YDB - DB Overview - IOPool, он показывает потребление CPU пулом управления блокирующими операциями ввода-вывода при операциях чтения-записи."
    new_description = get_description_for_graf(image_name, purpose_graf, purpose_test)
    replace_text(doc_name, "<LLM_IOPool>", new_description)

    # 18_TxReadOnly
    image_name = 'save/image18.png'
    purpose_graf = "Это график YDB - DB Overview - TxReadOnly, он показывает задержку для чтения на протяжении всего теста."
    new_description = get_description_for_graf(image_name, purpose_graf, purpose_test)
    replace_text(doc_name, "<LLM_TxReadOnly>", new_description)

    # 19_TxWriteOnly
    image_name = 'save/image19.png'
    purpose_graf = "Это график YDB - DB Overview - TxWriteOnly, он показывает задержку записи на протяжении всего теста."
    new_description = get_description_for_graf(image_name, purpose_graf, purpose_test)
    replace_text(doc_name, "<LLM_TxWriteOnly>", new_description)

    # 20_TxReadWrite
    image_name = 'save/image20.png'
    purpose_graf = "Это график YDB - DB Overview - TxReadWrite, он показывает задержку для чтения/записи на протяжении всего теста."
    new_description = get_description_for_graf(image_name, purpose_graf, purpose_test)
    replace_text(doc_name, "<LLM_TxReadWrite>", new_description)

    # 21_DataShardThroughput
    image_name = 'save/image21.png'
    purpose_graf = "Это график YDB - DB Overview - DataShardThroughput, он показывает количество строк, обрабатываемых в секунду."
    new_description = get_description_for_graf(image_name, purpose_graf, purpose_test)
    replace_text(doc_name, "<LLM_DataShardThroughput>", new_description)

    # 22_ShardDistribution
    image_name = 'save/image22.png'
    purpose_graf = "Это график YDB - DB Overview - ShardDistribution, он показывает распределения нагрузки между шардами."
    new_description = get_description_for_graf(image_name, purpose_graf, purpose_test)
    replace_text(doc_name, "<LLM_ShardDistribution>", new_description)

    # 23_Memory
    image_name = 'save/image23.png'
    purpose_graf = "Это график YDB - DB Overview - Memory, он показывает утилизацию памяти на каждом хосте БД."
    new_description = get_description_for_graf(image_name, purpose_graf, purpose_test)
    replace_text(doc_name, "<LLM_Memory>", new_description)

    # DB status #

    # 24_Disk
    image_name = 'save/image24.png'
    purpose_graf = "Это график YDB - DB Status - Disk, он показывает динамику использования дискового пространства БД."
    new_description = get_description_for_graf(image_name, purpose_graf, purpose_test)
    replace_text(doc_name, "<LLM_Disk>", new_description)

    print(f"25 descriptions for LLM were successfully replaced in doc")


# **Заменяем*описания*для*графиков*ответом*альфагена********************************************************************


# **Вставляем*данные*за*тест*в*таблицу*с*ошибками***********************************************************************

def insert_table_error_info(doc_name):
    time_from = config["start_time"]
    time_to = config["end_time"]

    data = get_content_from_panel("Error Info", time_from, time_to)

    # Создаем список для хранения агрегированных данных
    aggregated_data = []

    # Обрабатываем каждый фрейм
    for frame in data['results']['A']['frames']:
        # Извлекаем информацию о запросе/ошибке
        labels = frame['schema']['fields'][1]['labels']
        response_code = labels['responseCode']
        response_message = labels['responseMessage']
        transaction = labels.get('transaction', 'N/A')

        # Суммируем все значения
        total_count = sum(frame['data']['values'][1])

        # Добавляем в список
        aggregated_data.append({
            'response_code': response_code,
            'transaction': transaction,
            'response_message': response_message,
            'total_count': total_count
        })

    sorted_data = sorted(aggregated_data, key=lambda x: x['total_count'], reverse=True)[:5]

    doc = Document(doc_name)
    table = doc.tables[2]

    count_row_for_insert = len(sorted_data) if len(sorted_data) <= 5 else 5

    for row in range(count_row_for_insert):
        # получаем все объекты с ячейками в строке для вставки
        cells = table.rows[row + 1].cells
        cells[0].text = sorted_data[row]['transaction']
        cells[1].text = sorted_data[row]['response_code']
        cells[2].text = sorted_data[row]['response_message']
        cells[3].text = str(sorted_data[row]['total_count'])

    doc.save(doc_name)
    print(f"В документе {doc_name} сохранена таблица с ошибками")


# **Вставляем*данные*за*тест*в*таблицу*с*ошибками***********************************************************************

# **Делаем*замену*всех*картинок*в*документе*****************************************************************************

def replace_images(doc_name):
    count_photos = 24

    # print(f"\nStart to replace {count_photos} photos in the Doc \"{docName}\"")

    # Открываем новый DOCX-файл как архив
    with zipfile.ZipFile(doc_name, 'r') as docx:
        # Распаковываем содержимое во временную директорию
        tmp_dir = 'tmp'
        docx.extractall(tmp_dir)
        # Заменяем каждое изображение в документе
        for i in range(1, count_photos + 1):
            # Создаем новый файл изображения
            new_image_path = f"save/image{i}.png"
            # Находим путь к старому файлу изображения в директории media
            old_image_path = os.path.join(tmp_dir, 'word', 'media', f"image{i}.png")
            # Открываем исходное изображение с помощью Pillow
            with Image.open(new_image_path) as img:
                # Улучшаем качество изображения и сохраняем его во временную директорию
                img.save(old_image_path, quality=95)
            # Заменяем файл изображения

        # Упаковываем измененный документ обратно в DOCX-формат
        with zipfile.ZipFile(doc_name, 'w') as new_docx:
            for root, dirs, files in os.walk(tmp_dir):
                for file in files:
                    path = os.path.join(root, file)
                    rel_path = os.path.relpath(path, tmp_dir)
                    new_docx.write(path, rel_path)

    # Удаляем временные директории
    # TODO:УБРАТЬ ПОСЛЕ ОТЛАДКИ
    # shutil.rmtree(tmp_dir)
    print(f"{count_photos} photos were successfully replaced in doc \"{doc_name}\"\n")

# **Делаем*замену*всех*картинок*в*документе*****************************************************************************
